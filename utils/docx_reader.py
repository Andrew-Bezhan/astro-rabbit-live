"""
Модуль для чтения и конвертации файлов DOCX
"""

from docx import Document
import os


def read_docx_file(file_path):
    """
    Читает содержимое файла DOCX и возвращает текст
    
    Args:
        file_path (str): Путь к файлу DOCX
        
    Returns:
        str: Текстовое содержимое документа
    """
    try:
        doc = Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Пропускаем пустые параграфы
                text.append(paragraph.text)
        
        return '\n'.join(text)
    
    except Exception as e:
        print(f"Ошибка при чтении файла {file_path}: {e}")
        return None


def convert_docx_to_md(docx_path, md_path):
    """
    Конвертирует DOCX файл в Markdown
    
    Args:
        docx_path (str): Путь к исходному DOCX файлу
        md_path (str): Путь для сохранения MD файла
        
    Returns:
        bool: True если конвертация успешна, False в противном случае
    """
    try:
        doc = Document(docx_path)
        content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Определяем уровень заголовка по стилю
                style_name = paragraph.style.name if paragraph.style else None
                if style_name and 'Heading 1' in style_name:
                    content.append(f"# {text}")
                elif style_name and 'Heading 2' in style_name:
                    content.append(f"## {text}")
                elif style_name and 'Heading 3' in style_name:
                    content.append(f"### {text}")
                elif style_name and 'Heading 4' in style_name:
                    content.append(f"#### {text}")
                else:
                    content.append(text)
                content.append("")  # Пустая строка между параграфами
        
        # Сохраняем в файл
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        print(f"Файл успешно конвертирован: {md_path}")
        return True
        
    except Exception as e:
        print(f"Ошибка при конвертации файла: {e}")
        return False


def main():
    """Основная функция для тестирования"""
    docx_file = "План реализации проекта.docx"
    md_file = "План_реализации_проекта_readable.md"
    
    if os.path.exists(docx_file):
        print("Чтение содержимого DOCX файла...")
        content = read_docx_file(docx_file)
        
        if content:
            print("Первые 500 символов:")
            print(content[:500])
            print("\n" + "="*50 + "\n")
            
            print("Конвертация в Markdown...")
            if convert_docx_to_md(docx_file, md_file):
                print(f"Файл сохранен как: {md_file}")
            else:
                print("Ошибка при конвертации")
        else:
            print("Не удалось прочитать файл")
    else:
        print(f"Файл {docx_file} не найден")


if __name__ == "__main__":
    main()

